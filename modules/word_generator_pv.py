from io import BytesIO
from docx import Document
import pandas as pd

def _add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)

def _add_paragraph(doc, text):
    if text:
        doc.add_paragraph(text)
    else:
        doc.add_paragraph('')


def _add_table_from_df(doc, df, title=None):
    if title:
        doc.add_heading(title, level=2)
    if df is None or df.empty:
        doc.add_paragraph("(No hay datos)")
        return
    table = doc.add_table(rows=1, cols=len(df.columns))
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = str(col)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(df.columns):
            cells[i].text = str(row[col]) if not pd.isna(row[col]) else ''

def generar_pei_word(nombre_muni, tipo, mision, oei_df, aei_df, ruta, anexos):
    """Genera un documento Word (bytes) con la información suministrada."""
    doc = Document()

    doc.add_heading(f"Plan Estratégico Institucional - {nombre_muni}", level=0)
    doc.add_paragraph(f"Tipo de municipalidad: {tipo}")

    doc.add_heading("1. Misión", level=1)
    _add_paragraph(doc, mision)

    doc.add_heading("2. Objetivos Estratégicos Institucionales (OEI)", level=1)
    _add_table_from_df(doc, oei_df, title=None)

    doc.add_heading("3. Acciones Estratégicas Institucionales (AEI)", level=1)
    _add_table_from_df(doc, aei_df, title=None)

    doc.add_heading("4. Ruta Estratégica", level=1)
    _add_paragraph(doc, ruta)

    doc.add_heading("5. Anexo B-1", level=1)
    _add_paragraph(doc, anexos.get('B-1'))

    doc.add_heading("6. Anexo B-2", level=1)
    _add_paragraph(doc, anexos.get('B-2'))

    doc.add_heading("7. Anexo B-3", level=1)
    _add_paragraph(doc, anexos.get('B-3'))

    # Save to bytes
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()

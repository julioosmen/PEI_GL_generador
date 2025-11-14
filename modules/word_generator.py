from io import BytesIO
from docx import Document
import pandas as pd

# --- Utilidades ---
def _add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)

def _add_paragraph(doc, text):
    if text:
        doc.add_paragraph(str(text))
    else:
        doc.add_paragraph("(Sin contenido)")

def _add_table_from_df(doc, df, title=None):
    if title:
        doc.add_heading(title, level=2)

    if df is None or df.empty:
        doc.add_paragraph("(No hay datos para mostrar)")
        return

    table = doc.add_table(rows=1, cols=len(df.columns))
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = str(col)

    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(df.columns):
            val = row[col]
            cells[i].text = str(val) if not pd.isna(val) else ''

# --- Función principal ---
def generar_pei_word(nombre_muni, codigo, tipo, mision, oei_df, aei_df, ruta_df, anexo_b2_df, anexos=None):
    """
    Genera un documento Word con la información seleccionada:
    OEI, AEI, Ruta Estratégica (PGG) y Anexo B-2 (Políticas Nacionales).
    """
    doc = Document()
    doc.add_heading(f"Plan Estratégico Institucional - {codigo} {nombre_muni}", level=0)
    doc.add_paragraph(f"Tipo de municipalidad: {tipo}")

    # --- Misión ---
    _add_heading(doc, "1. Misión", level=1)
    _add_paragraph(doc, mision)

    # --- OEI ---
    _add_heading(doc, "2. Objetivos Estratégicos Institucionales (OEI)", level=1)
    _add_table_from_df(doc, oei_df)

    # --- AEI ---
    _add_heading(doc, "3. Acciones Estratégicas Institucionales (AEI)", level=1)
    _add_table_from_df(doc, aei_df)

    # --- Ruta Estratégica ---
    _add_heading(doc, "4. Ruta Estratégica (Vinculación con la PGG)", level=1)
    _add_table_from_df(doc, ruta_df)

    # --- Anexo B-2 ---
    _add_heading(doc, "5. Anexo B-2: Vinculación con Políticas Nacionales", level=1)
    _add_table_from_df(doc, anexo_b2_df)

    # --- Anexos adicionales ---
    if anexos:
        if 'B-1' in anexos:
            _add_heading(doc, "6. Anexo B-1", level=1)
            _add_paragraph(doc, anexos.get('B-1', ''))

        if 'B-3' in anexos:
            _add_heading(doc, "7. Anexo B-3", level=1)
            _add_paragraph(doc, anexos.get('B-3', ''))

    # --- Guardar a bytes ---
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()
